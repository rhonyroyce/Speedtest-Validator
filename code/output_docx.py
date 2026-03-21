"""Output Word document generator — creates RF_Throughput_Analysis.docx.

Generates a comprehensive RF analysis document with:
- Title page (site ID, date, analyst)
- Executive Summary (overall pass/fail, key findings)
- Site Configuration table (from CIQ)
- RF Parameter Deep Dive sections (per sector/tech)
- Summary tables (Band Config, Full Results, BLER Summary, BW vs Throughput, RX Chain)
- KPI Correlation Summary
- Glossary

Implementation: Claude Code Prompt 8 (Output Generation)
"""
import logging
from collections import defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

# RF/telecom glossary terms
GLOSSARY = {
    "RSRP": "Reference Signal Received Power — average power of resource elements carrying cell-specific reference signals (dBm).",
    "SINR": "Signal-to-Interference-plus-Noise Ratio — ratio of signal power to interference and noise (dB).",
    "RSRQ": "Reference Signal Received Quality — ratio of RSRP to total received wideband power (dB).",
    "PCI": "Physical Cell Identity — unique identifier (0-503) distinguishing cells on the same frequency.",
    "EARFCN": "E-UTRA Absolute Radio Frequency Channel Number — identifies the LTE carrier frequency.",
    "ARFCN": "Absolute Radio Frequency Channel Number — identifies the NR carrier frequency.",
    "EN-DC": "E-UTRA NR Dual Connectivity — LTE anchor with NR secondary carrier (NSA mode).",
    "NR-DC": "NR Dual Connectivity — dual NR carriers in SA mode.",
    "MIMO": "Multiple-Input Multiple-Output — multiple antenna paths for spatial multiplexing.",
    "SISO": "Single-Input Single-Output — single antenna path transmission.",
    "BLER": "Block Error Rate — percentage of transport blocks received with errors.",
    "MCS": "Modulation and Coding Scheme — index indicating modulation order and code rate.",
    "DAS": "Distributed Antenna System — network of antennas connected to a common source.",
    "CIQ": "Cell Information Questionnaire — engineering spreadsheet with cell configuration parameters.",
    "DL": "Downlink — data transmission from base station to user equipment.",
    "UL": "Uplink — data transmission from user equipment to base station.",
    "TX Power": "Transmit Power — power level at which the UE transmits to the base station (dBm).",
    "BW": "Bandwidth — width of the radio frequency channel (MHz).",
    "RRU": "Remote Radio Unit — radio equipment installed near the antenna.",
    "VSWR": "Voltage Standing Wave Ratio — measure of impedance matching in antenna system.",
    "AILG": "Automatic Idle-mode Load Guard — feature that adjusts RSSI targets.",
    "KPI": "Key Performance Indicator — measurable metric for network performance.",
}

# Alternating table row colors
ROW_COLOR_EVEN = "F2F2F2"
ROW_COLOR_ODD = "FFFFFF"


class OutputDocxGenerator:
    """Creates the RF_Throughput_Analysis.docx report."""

    def __init__(self, config: dict):
        self.config = config
        self.output_cfg = config.get("output", {})
        self.site_name = self.output_cfg.get("site_name", "UNKNOWN")

    def generate(
        self,
        results: list[dict],
        site_config: dict | list[dict],
        analysis_data: dict,
        output_path: str | Path,
    ) -> Path:
        """Create the complete Word document.

        Args:
            results: List of cell result dicts (same as output_xlsx).
            site_config: CIQ site configuration rows (sector, tech, band, BW, MIMO, PCI, EARFCN).
            analysis_data: Dict with aggregated analysis (observations, recommendations, kpi impacts).
            output_path: File path for the output document.

        Returns:
            Path to the written file.
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        self._apply_styles(doc)

        site_id = analysis_data.get("site_id", self.site_name)
        report_date = analysis_data.get("date", datetime.now().strftime("%Y-%m-%d"))

        self._add_title_page(doc, site_id, report_date)
        self._add_executive_summary(doc, results)
        self._add_site_config_table(doc, site_config)
        self._add_rf_deep_dive(doc, results)
        self._add_summary_tables(doc, results)
        self._add_kpi_correlation(doc, analysis_data)
        self._add_glossary(doc)

        doc.save(str(output_path))
        logger.info("Output DOCX written to %s", output_path)
        return output_path

    # ------------------------------------------------------------------
    # Section builders
    # ------------------------------------------------------------------

    def _add_title_page(self, doc: Document, site_id: str, date: str) -> None:
        """Add formatted title page."""
        # Add spacing before title
        for _ in range(6):
            doc.add_paragraph("")

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(f"RF Throughput Analysis \u2014 {site_id}")
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("DAS Validation Report")
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(f"Date: {date}")
        run.font.size = Pt(12)

        # Analyst
        analyst = self.output_cfg.get("analyst_name", "RF Validation Tool")
        analyst_para = doc.add_paragraph()
        analyst_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = analyst_para.add_run(f"Analyst: {analyst}")
        run.font.size = Pt(12)

        doc.add_page_break()

    def _add_executive_summary(self, doc: Document, results: list[dict]) -> None:
        """Add executive summary with pass/fail counts and key findings."""
        doc.add_heading("Executive Summary", level=1)

        # Count pass/fail
        total = len(results)
        pass_count = sum(
            1 for r in results
            if "PASS" in str(r.get("comment", "")).upper()
            and "FAIL" not in str(r.get("comment", "")).upper()
        )
        fail_count = total - pass_count

        # Overall status paragraph
        if fail_count == 0:
            status = "All tested cells passed validation thresholds."
        else:
            status = (
                f"{fail_count} of {total} tested cells failed one or more "
                f"validation thresholds and require investigation."
            )

        doc.add_paragraph(
            f"This report presents the RF throughput validation results for site "
            f"{self.site_name}. A total of {total} cell measurements were evaluated "
            f"against DAS validation thresholds. {status}"
        )

        # Key findings
        worst_rsrp = self._find_worst(results, "rsrp", min)
        worst_sinr = self._find_worst(results, "sinr", min)
        tp_failures = [
            r for r in results
            if "FAIL" in str(r.get("comment", "")).upper()
        ]

        findings = []
        if worst_rsrp is not None:
            findings.append(
                f"Worst RSRP: {worst_rsrp.get('rsrp')} dBm "
                f"({worst_rsrp.get('tech_sector', 'N/A')})"
            )
        if worst_sinr is not None:
            findings.append(
                f"Worst SINR: {worst_sinr.get('sinr')} dB "
                f"({worst_sinr.get('tech_sector', 'N/A')})"
            )
        if tp_failures:
            findings.append(
                f"{len(tp_failures)} cell(s) with throughput failures"
            )

        if findings:
            doc.add_paragraph(
                "Key findings: " + "; ".join(findings) + "."
            )

        # Pass/fail summary
        doc.add_paragraph(
            f"Overall result: {pass_count} PASS / {fail_count} FAIL "
            f"out of {total} cells tested."
        )

    def _add_site_config_table(self, doc: Document, site_config: dict | list) -> None:
        """Add site configuration table from CIQ data.

        Args:
            site_config: Either a list of cell dicts, or a dict keyed by sector
                         number from CIQReader.get_site_config_summary().
        """
        doc.add_heading("Site Configuration", level=1)
        doc.add_paragraph(
            "The following table summarizes the cell configuration parameters "
            "extracted from the CIQ engineering spreadsheet."
        )

        if not site_config:
            doc.add_paragraph("No CIQ configuration data available.")
            return

        # Flatten sector-grouped dict into a flat list of cell rows
        if isinstance(site_config, dict):
            flat_rows = []
            for sector, tech_groups in sorted(site_config.items()):
                for cell in tech_groups.get("lte", []):
                    flat_rows.append({
                        "sector": sector,
                        "technology": "LTE",
                        "band": cell.get("band", cell.get("dlChannelBandwidth", "")),
                        "bandwidth": cell.get("bandwidth_mhz", cell.get("dlChannelBandwidth", "")),
                        "mimo_config": "MIMO" if int(cell.get("noOfTxAntennas", 1)) >= 2 else "SISO",
                        "pci": cell.get("pci", cell.get("physicalLayerCellId", "")),
                        "earfcn": cell.get("earfcnDl", cell.get("earfcn", "")),
                    })
                for cell in tech_groups.get("nr", []):
                    flat_rows.append({
                        "sector": sector,
                        "technology": "NR",
                        "band": cell.get("band", cell.get("radioType", "")),
                        "bandwidth": cell.get("bandwidth_mhz", cell.get("channelBandwidth", "")),
                        "mimo_config": "MIMO" if int(cell.get("noOfTxAntennas", 1)) >= 2 else "SISO",
                        "pci": cell.get("pci", cell.get("nRPCI", "")),
                        "arfcn": cell.get("arfcnDl", cell.get("arfcn", "")),
                    })
            ciq_data = flat_rows
        else:
            ciq_data = site_config

        headers = ["Sector", "Technology", "Band", "Bandwidth", "MIMO Config", "PCI", "EARFCN/ARFCN"]
        table = doc.add_table(rows=1 + len(ciq_data), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        # Header row
        for col_idx, header in enumerate(headers):
            cell = table.rows[0].cells[col_idx]
            cell.text = header
            self._style_header_cell(cell)

        # Data rows
        for row_idx, row_data in enumerate(ciq_data, start=1):
            values = [
                str(row_data.get("sector", "")),
                str(row_data.get("technology", "")),
                str(row_data.get("band", "")),
                str(row_data.get("bandwidth", "")),
                str(row_data.get("mimo_config", "")),
                str(row_data.get("pci", "")),
                str(row_data.get("earfcn", row_data.get("arfcn", ""))),
            ]
            for col_idx, value in enumerate(values):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = value
                self._style_data_cell(cell, row_idx)

    def _add_rf_deep_dive(self, doc: Document, results: list[dict]) -> None:
        """Add per-sector/technology RF parameter analysis sections."""
        doc.add_heading("RF Parameter Analysis", level=1)

        # Group results by sector
        by_sector = defaultdict(list)
        for r in results:
            sector = r.get("tech_sector") or "Unknown"
            by_sector[sector].append(r)

        for sector, sector_results in sorted(by_sector.items()):
            doc.add_heading(sector, level=2)

            for cell_result in sector_results:
                conn_mode = cell_result.get("connection_mode", "")
                bw = cell_result.get("bandwidth", "")
                doc.add_heading(f"{conn_mode} \u2014 {bw}", level=3)

                # RF parameters table
                rf_table = doc.add_table(rows=5, cols=2)
                rf_table.style = "Table Grid"
                rf_params = [
                    ("RSRP", f"{cell_result.get('rsrp', 'N/A')} dBm"),
                    ("SINR", f"{cell_result.get('sinr', 'N/A')} dB"),
                    ("RSRQ", f"{cell_result.get('rsrq', 'N/A')} dB"),
                    ("UE TX Power", f"{cell_result.get('tx_power', 'N/A')} dBm"),
                ]
                rf_table.rows[0].cells[0].text = "Parameter"
                rf_table.rows[0].cells[1].text = "Value"
                self._style_header_cell(rf_table.rows[0].cells[0])
                self._style_header_cell(rf_table.rows[0].cells[1])
                for i, (param, val) in enumerate(rf_params, start=1):
                    rf_table.rows[i].cells[0].text = param
                    rf_table.rows[i].cells[1].text = str(val)

                doc.add_paragraph("")  # spacer

                # Throughput
                dl = cell_result.get("dl_throughput")
                ul = cell_result.get("ul_throughput")
                comment = cell_result.get("comment", "")
                doc.add_paragraph(
                    f"Throughput: DL {dl if dl is not None else 'N/A'} Mbps, "
                    f"UL {ul if ul is not None else 'N/A'} Mbps. "
                    f"Result: {comment}"
                )

                # Observations
                observations = cell_result.get("observations", "")
                if observations:
                    doc.add_paragraph(f"Observations: {observations}")

                # Recommendations
                recommendations = cell_result.get("recommendations", "")
                if recommendations:
                    doc.add_paragraph(f"Recommendations: {recommendations}")

    def _add_summary_tables(self, doc: Document, results: list[dict]) -> None:
        """Add summary tables: Band Config, Full Results."""
        doc.add_heading("Summary Tables", level=1)

        # --- Full Results Table ---
        doc.add_heading("Validation Results Summary", level=2)

        if not results:
            doc.add_paragraph("No results to display.")
            return

        summary_headers = [
            "Tech/Sector", "Connection Mode", "RSRP", "SINR",
            "DL (Mbps)", "UL (Mbps)", "Result",
        ]
        table = doc.add_table(rows=1 + len(results), cols=len(summary_headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        for col_idx, header in enumerate(summary_headers):
            cell = table.rows[0].cells[col_idx]
            cell.text = header
            self._style_header_cell(cell)

        for row_idx, r in enumerate(results, start=1):
            comment = str(r.get("comment", ""))
            result_label = "PASS" if ("PASS" in comment.upper() and "FAIL" not in comment.upper()) else "FAIL"
            values = [
                str(r.get("tech_sector", "")),
                str(r.get("connection_mode", "")),
                str(r.get("rsrp", "N/A")),
                str(r.get("sinr", "N/A")),
                str(r.get("dl_throughput", "N/A")),
                str(r.get("ul_throughput", "N/A")),
                result_label,
            ]
            for col_idx, value in enumerate(values):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = value
                self._style_data_cell(cell, row_idx)

        # --- Band Configuration Table ---
        doc.add_heading("Band Configuration Summary", level=2)

        # Aggregate unique band configurations
        band_configs = {}
        for r in results:
            key = (r.get("connection_mode", ""), r.get("bandwidth", ""))
            if key not in band_configs:
                band_configs[key] = {"count": 0, "pass": 0, "fail": 0}
            band_configs[key]["count"] += 1
            comment = str(r.get("comment", "")).upper()
            if "PASS" in comment and "FAIL" not in comment:
                band_configs[key]["pass"] += 1
            else:
                band_configs[key]["fail"] += 1

        bc_headers = ["Connection Mode", "Bandwidth", "Cells Tested", "Pass", "Fail"]
        bc_table = doc.add_table(rows=1 + len(band_configs), cols=len(bc_headers))
        bc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        bc_table.style = "Table Grid"

        for col_idx, header in enumerate(bc_headers):
            cell = bc_table.rows[0].cells[col_idx]
            cell.text = header
            self._style_header_cell(cell)

        for row_idx, ((conn_mode, bw), stats) in enumerate(sorted(band_configs.items()), start=1):
            values = [str(conn_mode), str(bw), str(stats["count"]), str(stats["pass"]), str(stats["fail"])]
            for col_idx, value in enumerate(values):
                cell = bc_table.rows[row_idx].cells[col_idx]
                cell.text = value
                self._style_data_cell(cell, row_idx)

    def _add_kpi_correlation(self, doc: Document, analysis_data: dict) -> None:
        """Add KPI correlation section linking RF findings to KPI impacts."""
        doc.add_heading("KPI Correlation Summary", level=1)

        doc.add_paragraph(
            "This section maps observed RF conditions to potential Key Performance "
            "Indicator impacts across the six MS2 KPI framework domains: "
            "Availability (AVL), Accessibility (ACC), Retainability (RET), "
            "Capacity (CAP), Mobility (MOB), and Power (PWR)."
        )

        # Per-cell KPI impacts
        kpi_impacts = analysis_data.get("kpi_impacts", [])
        if kpi_impacts:
            for impact in kpi_impacts:
                cell_label = impact.get("cell", "Unknown cell")
                doc.add_heading(cell_label, level=2)
                kpi_text = impact.get("kpi_impact", "No KPI impact data available.")
                doc.add_paragraph(kpi_text)
        else:
            doc.add_paragraph(
                "No significant KPI impacts were identified during this validation."
            )

    def _add_glossary(self, doc: Document) -> None:
        """Add RF/telecom terminology glossary."""
        doc.add_heading("Glossary", level=1)

        table = doc.add_table(rows=1 + len(GLOSSARY), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        # Header
        table.rows[0].cells[0].text = "Term"
        table.rows[0].cells[1].text = "Definition"
        self._style_header_cell(table.rows[0].cells[0])
        self._style_header_cell(table.rows[0].cells[1])

        # Entries (alphabetical)
        for row_idx, (term, definition) in enumerate(sorted(GLOSSARY.items()), start=1):
            table.rows[row_idx].cells[0].text = term
            run = table.rows[row_idx].cells[0].paragraphs[0].runs[0]
            run.font.bold = True
            table.rows[row_idx].cells[1].text = definition
            self._style_data_cell(table.rows[row_idx].cells[0], row_idx)
            self._style_data_cell(table.rows[row_idx].cells[1], row_idx)

    # ------------------------------------------------------------------
    # Styling
    # ------------------------------------------------------------------

    def _apply_styles(self, doc: Document) -> None:
        """Configure document-wide styles: fonts, headings, margins."""
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(6)

        # Heading styles
        for level, (size, color) in enumerate(
            [(16, "1F497D"), (14, "2E75B6"), (12, "4472C4")], start=1
        ):
            heading_style = doc.styles[f"Heading {level}"]
            heading_style.font.name = "Calibri"
            heading_style.font.size = Pt(size)
            heading_style.font.bold = True
            heading_style.font.color.rgb = RGBColor(
                int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
            )

        # Set default section margins
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

    @staticmethod
    def _style_header_cell(cell) -> None:
        """Apply header styling to a table cell."""
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Background color
        shading = cell._element.get_or_add_tcPr()
        shading_elem = shading.makeelement(
            qn("w:shd"),
            {qn("w:fill"): "4472C4", qn("w:val"): "clear"},
        )
        shading.append(shading_elem)

    @staticmethod
    def _style_data_cell(cell, row_idx: int) -> None:
        """Apply alternating row color to a data cell."""
        color = ROW_COLOR_EVEN if row_idx % 2 == 0 else ROW_COLOR_ODD
        shading = cell._element.get_or_add_tcPr()
        shading_elem = shading.makeelement(
            qn("w:shd"),
            {qn("w:fill"): color, qn("w:val"): "clear"},
        )
        shading.append(shading_elem)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _find_worst(results: list[dict], key: str, func) -> dict | None:
        """Find the result with the worst (min/max) value for a given key."""
        valid = [r for r in results if isinstance(r.get(key), (int, float))]
        if not valid:
            return None
        return func(valid, key=lambda r: r[key])
